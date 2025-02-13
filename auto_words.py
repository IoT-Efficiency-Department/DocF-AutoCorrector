#需要安装python-docx包

import docx
from docx import Document
from docx.shared import Pt,RGBColor,Cm,Inches,Length # 字号，设置像素、缩进,颜色,宽度，厘米，英寸等
from docx.oxml.ns import qn # 中文字体
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_TAB_ALIGNMENT,WD_TAB_LEADER  #设置对象居中、对齐、制表符等


from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

文件 = Document('./testin.docx')

封面页数 = 1
摘要页数 = 0
目录页数 = 0

page = 1

def create_element(name):
    return OxmlElement(name)
def create_attribute(element, name, value):
    element.set(qn(name), value)
def add_page_number(run):
    # 创建页码开始字段
    # 页码开始
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    # 创建页码指令文本
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    # 创建页码结束字段
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    # 将它们添加到运行对象中
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)



j = int(len(文件.sections))
print(j)
for i in range(j):
    节 = 文件.sections[i]


#确定封面节设定正确
    # 假设封面是第一节，设置第一页无页码
    if 封面页数 > 0 :
        if i == 0 :
            封面页 = 文件.sections[i]
            封面页脚 = 封面页.footer
            封面页眉 = 封面页.header
            # 封面页脚段落 = 封面页脚.add_paragraph()
            for paragraph in 封面页脚.paragraphs:
                paragraph.clear()  # 清空封面页脚内容
            for paragraph in 封面页眉.paragraphs:
                paragraph.clear()  # 清空封面页眉内容
    if 摘要页数 > 0:
        if i == 封面页数:
            摘要页 = 文件.sections[i]
            摘要页脚 = 摘要页.footer
            摘要页眉 = 摘要页.header
            for paragraph in 摘要页脚.paragraphs:
                paragraph.clear()  # 清空封面页脚内容
            for paragraph in 摘要页眉.paragraphs:
                paragraph.clear()  # 清空封面页眉内容
    if 目录页数 > 0:
        if i == 封面页数 + 摘要页数 :

            for n in 目录页数:
                目录页 = 文件.sections[i+n]
                目录页脚 = 目录页.footer
                目录页眉 = 目录页.header

                for paragraph in 目录页脚.paragraphs:
                    paragraph.clear()  # 清空目录页脚内容
                    paragraph.text = str(i + n + 1)  # 设置目录页脚内容页码
                for paragraph in 目录页眉.paragraphs:
                    paragraph.clear()  # 清空目录页眉内容


    if 节.page_width == Cm(21) & 节.page_height == Cm(29.7): #页面高度
        print("页面正确！");
    else:
        # print("页面错误！")
        节.page_width = Cm(21)  # 页面宽度
        节.page_height = Cm(29.7)  # 页面高度

    节.top_margin = Cm(2.5) #页边距上
    节.bottom_margin = Cm(2.5) #页边距下
    节.left_margin = Cm(3) #页边距左
    节.right_margin = Cm(2) #页边距右
    节.header_distance = Cm(1.5) #页眉距离1.5厘米
    节.footer_distance = Cm(1.75) #页脚距离1.75厘米
    页眉 = 节.header
    页眉段落 =页眉.paragraphs[0]
    页眉段落数量 = len(页眉.paragraphs)

    # 删除页眉多余的段落
    if 页眉段落数量 > 1:
        for _ in range(页眉段落数量 - 1):
            页眉.paragraphs[-1]._element.getparent().remove(页眉.paragraphs[-1]._element)

    for 块 in 页眉段落.runs:
        块.font.size = Pt(9)  # 页眉小五
        块.font.name = 'Times New Roman'
        块._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    页眉段落.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #设置页眉段落为居中

    页脚 = 节.footer



    页脚段落数量 = len(页脚.paragraphs)

    # 删除页脚多余的段落
    if 页脚段落数量 > 1:
        for _ in range(页脚段落数量 - 1):
            页脚.paragraphs[-1]._element.getparent().remove(页脚.paragraphs[-1]._element)

    #
    # 页脚段落 = 页脚.paragraphs[0]
    # for 块 in 页脚段落.runs:
    #     块.font.size = Pt(9)  # 页脚小五
    #     块.font.name = 'Times New Roman'
    #     块._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    #
    # 页脚段落.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 页脚 = 节.footer
    # paragraph = 页脚.paragraphs[0]  # 获取页脚的第一个段落
    # 页脚段落.clear()  # 清空现有内容


# add_page_number(文件.sections[0].footer.paragraphs[0].add_run())
# 文件.sections[0].footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 遍历每一节并添加页码
start_section_index = 1  # 从第2节开始添加页码（索引从0开始）
for i, section in enumerate(文件.sections):
    if i >= start_section_index:
        页脚 = section.footer
        页脚段落 = 页脚.paragraphs[0] if 页脚.paragraphs else 页脚.add_paragraph()
        页脚段落.clear()  # 清空现有内容
        add_page_number(页脚段落.add_run())
        页脚段落.alignment = WD_ALIGN_PARAGRAPH.CENTER


# 全文修改
for 段落 in 文件.paragraphs:
    for 块 in 段落.runs:
        # 一级标题
        if 段落.style.name == 'Heading 1':
            块.font.size = Pt(16)#三号
            块.font.bold = True  # 字体加粗

            块.font.name = 'Times New Roman'
            块._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

            块.font.color.rgb = RGBColor(0, 0, 0, )  # 颜色
            段落.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左顶格
            段落.paragraph_format.line_spacing = 1.5 #1.5倍行距
            段落.paragraph_format.space_before = Pt(6) #段前
            段落.paragraph_format.space_after = Pt(6) #段后
            段落.paragraph_format.left_indent = Inches(0)  # 正文前
            段落.paragraph_format.right_indent = Inches(0)  # 正文后
        elif 段落.style.name == 'Heading 2':
            块.font.size = Pt(15)# 小三号
            块.font.bold = True #字体加粗
            块.font.name = 'Times New Roman'
            块._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            块.font.color.rgb = RGBColor(0, 0, 0, )  # 颜色
            段落.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT # 左顶格
            段落.paragraph_format.line_spacing = 1.5  # 1.5倍行距
            段落.paragraph_format.space_before = Pt(6) #段前
            段落.paragraph_format.space_after = Pt(0) #段后
            段落.paragraph_format.left_indent = Inches(0)  # 正文前
            段落.paragraph_format.right_indent = Inches(0)  # 正文后
        elif 段落.style.name == 'Heading 3':
            块.font.size = Pt(14)# 小四
            块.font.bold = True  # 字体加粗
            块.font.name = 'Times New Roman'
            块._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            块.font.color.rgb = RGBColor(0, 0, 0, )  # 颜色
            段落.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            段落.paragraph_format.line_spacing = 1.5  # 1.5倍行距
            段落.paragraph_format.space_before = Pt(6) #段前
            段落.paragraph_format.space_after = Pt(0) #段后
            段落.paragraph_format.left_indent = Inches(0)  # 正文前
            段落.paragraph_format.right_indent = Inches(0)  # 正文后
        elif 段落.style.name == 'Heading 4':
            块.font.size = Pt(12)#小四号
            块.font.bold = True  # 字体加粗
            块.font.name = 'Times New Roman'
            块._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            块.font.color.rgb = RGBColor(0, 0, 0, )  # 颜色
            段落.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            段落.paragraph_format.line_spacing = 1.5  # 1.5倍行距
            段落.paragraph_format.space_before = Pt(0)  # 段前
            段落.paragraph_format.space_after = Pt(0)  # 段后
            段落.paragraph_format.left_indent = Inches(0)  # 正文前
            段落.paragraph_format.right_indent = Inches(0)  # 正文后
        elif 段落.style.name == 'Heading 5':
            块.font.size = Pt(12)#小四号
            块.font.bold = True  # 字体加粗
            块.font.name = 'Times New Roman'
            块._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            块.font.color.rgb = RGBColor(0, 0, 0, )  # 颜色
            段落.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            段落.paragraph_format.line_spacing = 1.5  # 1.5倍行距
            段落.paragraph_format.space_before = Pt(0)  # 段前
            段落.paragraph_format.space_after = Pt(0)  # 段后
            段落.paragraph_format.left_indent = Inches(0)  # 正文前
            段落.paragraph_format.right_indent = Inches(0)  # 正文后
        elif 段落.style.name == 'Heading 6':
            块.font.size = Pt(12)#小四号
            块.font.bold = True  # 字体加粗
            块.font.name = 'Times New Roman'
            块._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            块.font.color.rgb = RGBColor(0, 0, 0, )  # 颜色
            段落.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            段落.paragraph_format.line_spacing = 1.5  # 1.5倍行距
            段落.paragraph_format.space_before = Pt(0)  # 段前
            段落.paragraph_format.space_after = Pt(0)  # 段后
            段落.paragraph_format.left_indent = Inches(0)  # 正文前
            段落.paragraph_format.right_indent = Inches(0)  # 正文后
        #正文
        elif 段落.style.name == 'Normal':
            块.font.size = Pt(12)#小四号
            块.font.name = 'Times New Roman' # 设置英文字体
            块._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
            块.font.color.rgb = RGBColor(0,0,0,) #颜色
            段落.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐
            段落.paragraph_format.line_spacing = 1.5  # 1.5倍行距
            段落.paragraph_format.space_before = Pt(0) #段前
            段落.paragraph_format.space_after = Pt(0) #段后
            段落.paragraph_format.left_indent = Inches(0) #正文前
            段落.paragraph_format.right_indent = Inches(0) #正文后
            段落.paragraph_format.first_line_indent = Cm(0.823) #首行缩进

文件.save('./testout.docx')